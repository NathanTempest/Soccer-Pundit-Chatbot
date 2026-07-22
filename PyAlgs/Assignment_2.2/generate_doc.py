import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_page_number(run):
    """Inserts a page number field into a run object."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_title_page_paragraph(doc, text, bold=False, space_before=Pt(0), space_after=Pt(0)):
    """Helper to add centered, double-spaced paragraph for the title page."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading_1(doc, text):
    """Helper to add APA Heading 1 (centered, bold, title case)."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_heading_2(doc, text):
    """Helper to add APA Heading 2 (left-aligned, bold, title case)."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_body_paragraph(doc, text):
    """Helper to add APA standard body paragraph (double-spaced, 0.5 in indent)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_code_paragraph(doc, text):
    """Helper to add a code block (single-spaced, left-indented, Courier New)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p

def add_output_paragraph(doc, text):
    """Helper to add a console output block (single-spaced, left-indented, italics, Courier New)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(100, 100, 100)
    return p

def add_placeholder(doc, label):
    """Helper to add a prominent, styled placeholder for screenshot insertion."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"\n[--- INSERT SCREENSHOT FOR {label.upper()} HERE ---]\n")
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(180, 0, 0) # Dark red for high visibility
    return p

def add_reference_paragraph(doc, text):
    """Helper to add an APA reference with a hanging indent (0.5 in)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def main():
    doc = Document()
    
    # 1. Page Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Configure Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        hp.paragraph_format.line_spacing = 1.0
        hp.paragraph_format.space_after = Pt(0)
        hp.paragraph_format.space_before = Pt(0)
        run = hp.add_run(" ")
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        add_page_number(run)

    # 2. TITLE PAGE (APA 7th Student Standard)
    add_title_page_paragraph(doc, "", space_before=Pt(72))
    add_title_page_paragraph(doc, "Module 2 Workshop: Sequences, Functions, and Descriptive Statistics", bold=True, space_after=Pt(24))
    add_title_page_paragraph(doc, "Nathaniel Jason")
    add_title_page_paragraph(doc, "College of Science, Engineering, and Technology, Grand Canyon University")
    add_title_page_paragraph(doc, "CST-111: Introduction to Computer Science and Information Technology")
    add_title_page_paragraph(doc, "Dr. Instructor")
    add_title_page_paragraph(doc, "July 20, 2026")
    
    doc.add_page_break()

    # 3. BODY OF WORK
    add_heading_1(doc, "Module 2 Workshop: Sequences, Functions, and Descriptive Statistics")
    
    intro_p1 = (
        "Sequences form the foundational structural elements of Python programming, enabling developers "
        "to organize, store, and manipulate data collections systematically. In Python, sequences are "
        "categorized by their properties, with lists and tuples being the most widely utilized. A list is "
        "a mutable sequence, allowing for dynamic additions, deletions, and alterations of its contents. "
        "This mutability makes lists ideal for managing datasets that change over time, such as transaction "
        "logs, user inputs, or sensory data streams. Conversely, a tuple is an immutable sequence, meaning "
        "its elements cannot be altered, added, or removed once initialized. This immutability ensures data "
        "integrity, making tuples suitable for fixed datasets like geographical coordinates, database keys, "
        "or configuration parameters."
    )
    add_body_paragraph(doc, intro_p1)
    
    intro_p2 = (
        "Beyond storage, the manipulation of sequences through indexing, slicing, and built-in functions "
        "represents a critical aspect of program development. Slicing provides a clean and expressive "
        "mechanism to retrieve subsegments of data, while built-in functions like filter() coupled with "
        "lambda expressions enable functional programming patterns to extract relevant records efficiently. "
        "Furthermore, statistical analysis is intrinsically tied to data structures; lists serve as the "
        "inputs to functions that compute measures of central tendency (mean, median, mode) and dispersion "
        "(variance, standard deviation). This report presents the design, pseudocodes, code implementations, "
        "and empirical outputs for five programming exercises demonstrating these programming paradigms."
    )
    add_body_paragraph(doc, intro_p2)

    # --- Exercise 1 ---
    add_heading_2(doc, "Exercise 1: Datetime Function")
    ex1_narrative = (
        "The first exercise demonstrates the use of standard library integration, specifically utilizing Python's "
        "datetime module. A parameterless function named date_and_time() is defined, containing a statement that "
        "calls the today() method on the datetime class. Calling this function outputs the current local date "
        "and time as an object representing year, month, day, hours, minutes, seconds, and microseconds."
    )
    add_body_paragraph(doc, ex1_narrative)
    
    ex1_pseudocode = (
        "Pseudocode:\n"
        "1. IMPORT the datetime module.\n"
        "2. DEFINE function date_and_time():\n"
        "    a. CALL datetime.datetime.today() and print the result.\n"
        "3. CALL the date_and_time() function."
    )
    add_body_paragraph(doc, ex1_pseudocode)
    
    ex1_code = (
        "import datetime\n\n"
        "def date_and_time():\n"
        "    \"\"\"Prints the current date and time using datetime.datetime.today().\"\"\"\n"
        "    print(datetime.datetime.today())\n\n"
        "if __name__ == '__main__':\n"
        "    date_and_time()"
    )
    add_code_paragraph(doc, ex1_code)
    
    ex1_output = (
        "Console Output:\n"
        "2026-07-19 22:37:24.614606"
    )
    add_output_paragraph(doc, ex1_output)
    add_placeholder(doc, "Exercise 1 Execution Screenshot")

    # --- Exercise 2 ---
    add_heading_2(doc, "Exercise 2: Temperature Conversion Chart")
    ex2_narrative = (
        "This exercise explores program development using arithmetic functions and tabular output alignment. "
        "A fahrenheit(celsius) function is implemented to return the Fahrenheit equivalent of a Celsius temperature "
        "using the formula F = (9/5) * C + 32. A for-loop is then utilized to generate Celsius values from 0 to 100. "
        "For each value, the fahrenheit() function is called, and the resulting Celsius-Fahrenheit pairs are printed "
        "in a neat, right-aligned tabular structure formatted with one decimal place of precision."
    )
    add_body_paragraph(doc, ex2_narrative)
    
    ex2_pseudocode = (
        "Pseudocode:\n"
        "1. DEFINE function fahrenheit(celsius):\n"
        "    a. RETURN (9 / 5) * celsius + 32.\n"
        "2. PRINT headers 'Celsius' and 'Fahrenheit' aligned.\n"
        "3. PRINT separator line.\n"
        "4. FOR c from 0 to 100 (inclusive):\n"
        "    a. COMPUTE f = fahrenheit(c).\n"
        "    b. PRINT c and f right-aligned, formatting f with 1 decimal digit."
    )
    add_body_paragraph(doc, ex2_pseudocode)
    
    ex2_code = (
        "def fahrenheit(celsius):\n"
        "    \"\"\"Returns the Fahrenheit equivalent of a Celsius temperature.\"\"\"\n"
        "    return (9 / 5) * celsius + 32\n\n"
        "def main():\n"
        "    print(f\"{'Celsius':>7}  {'Fahrenheit':>11}\")\n"
        "    print('-' * 20)\n"
        "    for c in range(101):\n"
        "        f = fahrenheit(c)\n"
        "        print(f\"{c:>7}  {f:>11.1f}\")\n\n"
        "if __name__ == '__main__':\n"
        "    main()"
    )
    add_code_paragraph(doc, ex2_code)
    
    ex2_output = (
        "Console Output:\n"
        "Celsius   Fahrenheit\n"
        "--------------------\n"
        "      0         32.0\n"
        "      1         33.8\n"
        "      2         35.6\n"
        "     ...\n"
        "    100        212.0"
    )
    add_output_paragraph(doc, ex2_output)
    add_placeholder(doc, "Exercise 2 Execution Tabular Output Screenshot")

    # --- Exercise 3 ---
    add_heading_2(doc, "Exercise 3: String Slicing Operations")
    ex3_narrative = (
        "Slicing is a core technique in Python that allows parts of a sequence to be extracted without modifying "
        "the original object. In this exercise, the string 'abcdefghijklmnopqrstuvwxyz' is sliced in seven distinct "
        "ways: extracting the first half with start and end indices, using only the end index, extracting the "
        "second half with start and end indices, using only the start index, retrieving every second letter starting "
        "with 'a', reversing the entire string, and extracting every third letter in reverse starting with 'z'."
    )
    add_body_paragraph(doc, ex3_narrative)
    
    ex3_pseudocode = (
        "Pseudocode:\n"
        "1. SET alphabet = 'abcdefghijklmnopqrstuvwxyz'.\n"
        "2. COMPUTE and PRINT separate slices:\n"
        "    a. First half using start/end: alphabet[0:13]\n"
        "    b. First half using only end: alphabet[:13]\n"
        "    c. Second half using start/end: alphabet[13:26]\n"
        "    d. Second half using only start: alphabet[13:]\n"
        "    e. Every second letter starting with 'a': alphabet[::2]\n"
        "    f. Entire string in reverse: alphabet[::-1]\n"
        "    g. Every third letter in reverse starting with 'z': alphabet[::-3]"
    )
    add_body_paragraph(doc, ex3_pseudocode)
    
    ex3_code = (
        "def main():\n"
        "    alphabet = 'abcdefghijklmnopqrstuvwxyz'\n"
        "    print(f\"a. {alphabet[0:13]}\")\n"
        "    print(f\"b. {alphabet[:13]}\")\n"
        "    print(f\"c. {alphabet[13:26]}\")\n"
        "    print(f\"d. {alphabet[13:]}\")\n"
        "    print(f\"e. {alphabet[::2]}\")\n"
        "    print(f\"f. {alphabet[::-1]}\")\n"
        "    print(f\"g. {alphabet[::-3]}\")\n\n"
        "if __name__ == '__main__':\n"
        "    main()"
    )
    add_code_paragraph(doc, ex3_code)
    
    ex3_output = (
        "Console Output:\n"
        "a. abcdefghijklm\n"
        "b. abcdefghijklm\n"
        "c. nopqrstuvwxyz\n"
        "d. nopqrstuvwxyz\n"
        "e. acegikmoqsuwy\n"
        "f. zyxwvutsrqponmlkjihgfedcba\n"
        "g. zwtqnkheb"
    )
    add_output_paragraph(doc, ex3_output)
    add_placeholder(doc, "Exercise 3 Execution Slicing Screenshot")

    # --- Exercise 4 ---
    add_heading_2(doc, "Exercise 4: Finding the People with a Specified Last Name")
    ex4_narrative = (
        "This exercise demonstrates structural filtering on collections of complex objects. A list of tuples is "
        "created, with each tuple containing a first and a last name. The built-in filter() function is called with "
        "a lambda expression. The lambda expression inspects the second element (index 1) of each tuple to check "
        "if it matches the string 'Jones'. The resulting iterator is iterated over to output the filtered tuples."
    )
    add_body_paragraph(doc, ex4_narrative)
    
    ex4_pseudocode = (
        "Pseudocode:\n"
        "1. CREATE a list of tuples containing first and last names.\n"
        "2. FILTER the list using a lambda function: target last name equals 'Jones'.\n"
        "3. PRINT header 'People with last name Jones'.\n"
        "4. FOR each matching tuple in the filtered iterator:\n"
        "    a. PRINT the tuple."
    )
    add_body_paragraph(doc, ex4_pseudocode)
    
    ex4_code = (
        "def main():\n"
        "    people = [\n"
        "        ('Alice', 'Jones'), ('Bob', 'Smith'), ('Carol', 'Jones'),\n"
        "        ('David', 'Brown'), ('Eve', 'Jones'), ('Frank', 'Davis'),\n"
        "        ('Grace', 'Jones'), ('Hank', 'Wilson')\n"
        "    ]\n"
        "    jones_people = filter(lambda person: person[1] == 'Jones', people)\n"
        "    print(\"People with last name 'Jones':\")\n"
        "    for person in jones_people:\n"
        "        print(person)\n\n"
        "if __name__ == '__main__':\n"
        "    main()"
    )
    add_code_paragraph(doc, ex4_code)
    
    ex4_output = (
        "Console Output:\n"
        "People with last name 'Jones':\n"
        "('Alice', 'Jones')\n"
        "('Carol', 'Jones')\n"
        "('Eve', 'Jones')\n"
        "('Grace', 'Jones')"
    )
    add_output_paragraph(doc, ex4_output)
    add_placeholder(doc, "Exercise 4 Execution Filtering Screenshot")

    # --- Exercise 5 ---
    add_heading_2(doc, "Exercise 5: Survey Response Statistics")
    ex5_narrative = (
        "This exercise performs descriptive statistical analysis on a dataset of 20 survey responses rating cafeteria "
        "food on a scale of 1 to 5. Rating frequencies are computed using the list count() method. Descriptive "
        "statistics—minimum, maximum, range, mean, median, mode, variance, and standard deviation—are calculated "
        "using built-in functions, the standard 'statistics' library, and 'NumPy'."
    )
    add_body_paragraph(doc, ex5_narrative)
    
    ex5_nuance = (
        "Mathematical Note on Sample Output Discrepancy:\n"
        "A rigorous review of the textbook's sample output shows a variance of 1.293421052631579 and standard deviation "
        "of 1.1373020273294305 for a mean of 3.05. However, the sum of frequencies in the sample output (3+3+8+2+3) is 19. "
        "A sample variance of exactly 1.293421052631579 (which is 24.575/19) is mathematically impossible for a list of 20 "
        "integers. For the sum of squared differences to equal 24.575, the sum of squares of the integers would need to be "
        "210.625 (a non-integer). This indicates that the sample output either has a calculation typo or was generated "
        "from a slightly different dataset. When executed dynamically on the provided list of 20 ratings, the true mean "
        "is 2.9, the median is 3.0, the mode is 3, the sample variance is 1.568421052631579, and the sample standard "
        "deviation is 1.2523661815266247."
    )
    add_body_paragraph(doc, ex5_nuance)
    
    ex5_pseudocode = (
        "Pseudocode:\n"
        "1. INITIALIZE a list responses = [1, 2, 5, 4, 3, 5, 2, 1, 3, 3, 1, 4, 3, 3, 3, 2, 3, 3, 2, 5].\n"
        "2. PRINT 'Rating Frequencies:'.\n"
        "3. FOR each rating from 1 to 5:\n"
        "    a. COUNT occurrences in responses and print.\n"
        "4. COMPUTE stats:\n"
        "    a. Minimum = np.min(responses), Maximum = np.max(responses)\n"
        "    b. Range = Maximum - Minimum\n"
        "    c. Mean = statistics.mean(responses), Median = statistics.median(responses)\n"
        "    d. Mode = statistics.mode(responses)\n"
        "    e. Variance = np.var(responses, ddof=1) (sample variance)\n"
        "    f. Standard Deviation = np.std(responses, ddof=1) (sample std)\n"
        "5. PRINT calculated stats."
    )
    add_body_paragraph(doc, ex5_pseudocode)
    
    ex5_code = (
        "import statistics\n"
        "import numpy as np\n\n"
        "def main():\n"
        "    responses = [1, 2, 5, 4, 3, 5, 2, 1, 3, 3, 1, 4, 3, 3, 3, 2, 3, 3, 2, 5]\n"
        "    print(\"Rating Frequencies:\")\n"
        "    for rating in range(1, 6):\n"
        "        frequency = responses.count(rating)\n"
        "        print(f\"{rating}: {frequency}\")\n"
        "    print()\n"
        "    minimum = np.min(responses)\n"
        "    maximum = np.max(responses)\n"
        "    stat_range = maximum - minimum\n"
        "    mean = statistics.mean(responses)\n"
        "    median = statistics.median(responses)\n"
        "    mode = statistics.mode(responses)\n"
        "    variance = np.var(responses, ddof=1)\n"
        "    std_dev = np.std(responses, ddof=1)\n"
        "    print(\"Response Statistics:\")\n"
        "    print(f\"Minimum: {minimum}\")\n"
        "    print(f\"Maximum: {maximum}\")\n"
        "    print(f\"Range: {stat_range}\")\n"
        "    print(f\"Mean: {mean}\")\n"
        "    print(f\"Median: {float(median)}\")\n"
        "    print(f\"Mode: {mode}\")\n"
        "    print(f\"Variance: {variance}\")\n"
        "    print(f\"Standard Deviation: {std_dev}\")\n\n"
        "if __name__ == '__main__':\n"
        "    main()"
    )
    add_code_paragraph(doc, ex5_code)
    
    ex5_output = (
        "Console Output:\n"
        "Rating Frequencies:\n"
        "1: 3\n"
        "2: 4\n"
        "3: 8\n"
        "4: 2\n"
        "5: 3\n\n"
        "Response Statistics:\n"
        "Minimum: 1\n"
        "Maximum: 5\n"
        "Range: 4\n"
        "Mean: 2.9\n"
        "Median: 3.0\n"
        "Mode: 3\n"
        "Variance: 1.568421052631579\n"
        "Standard Deviation: 1.2523661815266247"
    )
    add_output_paragraph(doc, ex5_output)
    add_placeholder(doc, "Exercise 5 Execution Statistics Screenshot")

    # --- Reflection Section ---
    add_heading_1(doc, "Conceptual and Faith-Integrated Reflection")
    
    ref_p1 = (
        "The comparison between lists and tuples in Python relates directly to programmatic performance "
        "and data integrity. Lists are implemented as dynamic arrays, which allocate extra memory buffer "
        "to handle resizing operations efficiently. This mutability makes them flexible but introduces memory "
        "overhead. Conversely, tuples are static and immutable, allowing Python to allocate the exact memory "
        "required. As a result, tuples are faster to create, use less memory, and provide safety by preventing "
        "unintended modifications. They are often used as keys in dictionaries or returned from functions to "
        "ensure read-only integrity. Understanding these design trade-offs allows developers to build robust, "
        "performance-optimized applications."
    )
    add_body_paragraph(doc, ref_p1)
    
    ref_p2 = (
        "The use of advanced operations such as slicing and filter() provides elegant, functional programming "
        "mechanisms. Slicing syntax ([start:stop:step]) allows for quick subsegment extraction at highly optimized "
        "C-level speeds, avoiding the overhead of manual loops. The filter() function, paired with lambda functions, "
        "constructs memory-efficient iterators that evaluate predicates lazily, processing data elements only "
        "when requested. This functional approach results in readable and maintainable source code, which is "
        "essential when working with large, complex datasets."
    )
    add_body_paragraph(doc, ref_p2)
    
    ref_p3 = (
        "From a faith-integrated perspective, the analysis of data and programming is a reflection of stewardship "
        "and truth. Humanity was created in the likeness and image of God, endowed with intellect, reasoning, "
        "and the capability to solve complex problems. Applying programming and data analysis is an extension of "
        "honing these talents to serve others and make a positive difference in the world. Accurate computation and "
        "honest analysis of statistics represent a commitment to truth and integrity, echoing Proverbs 11:1, "
        "which states that 'a false balance is an abomination to the Lord, but a just weight is his delight.' By "
        "performing rigorous analysis, identifying mathematical inconsistencies, and structuring data into "
        "knowledge, programmers act as good stewards of information, aligning their skills with God's calling "
        "for order, honesty, and service."
    )
    add_body_paragraph(doc, ref_p3)

    # --- References ---
    doc.add_page_break()
    add_heading_1(doc, "References")
    
    ref1 = (
        "Deitel, P., & Deitel, H. (2020). Intro to Python for Computer Science and Data Science: "
        "Learning to Program with AI, Big Data and the Cloud. Pearson Education."
    )
    add_reference_paragraph(doc, ref1)
    
    ref2 = (
        "Python Software Foundation. (2026). The Python Language Reference, version 3.12. "
        "Available at https://www.python.org/"
    )
    add_reference_paragraph(doc, ref2)

    # Save Document relative to script directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    doc_path = os.path.join(script_dir, "Jason_Nathaniel_Assignment_2_2.docx")
    doc.save(doc_path)
    print(f"Document successfully created at: {os.path.abspath(doc_path)}")

if __name__ == '__main__':
    main()
