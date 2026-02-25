from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Title ---
title = doc.add_paragraph()
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_run = title.add_run("Nathaniel Jason")
title_run.bold = True
title_run.font.size = Pt(16)

# --- Contact Info ---
contact = doc.add_paragraph()
contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact.add_run("nathanjason2077@gmail.com | (716) 463-9623 | linkedin.com/in/natejason")

# --- Helper: Add Main Heading ---
def add_heading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.paragraph_format.space_before = Pt(10)

# --- Professional Summary ---
add_heading("Professional Summary")

summary_points = [
    "7+ years of professional experience in full stack and cloud development, specializing in AWS serverless solutions with Node.js, TypeScript, and Java.",
    "Proven expertise in designing, developing, and deploying microservices and event-driven architectures using AWS Lambda, API Gateway, SQS, SNS, and Step Functions.",
    "Strong experience with AWS CDK for infrastructure-as-code, enabling scalable and repeatable deployments.",
    "Extensive background in relational and NoSQL databases including Postgres, MySQL, DynamoDB, and MongoDB.",
    "Hands-on experience with CI/CD pipelines using GitLab and Jenkins, implementing automated testing and deployment workflows.",
    "Developed high-performance RESTful and GraphQL APIs with Node.js, Spring Boot, and Python frameworks.",
    "Implemented enterprise-grade monitoring and observability using CloudWatch, ELK Stack, and Prometheus.",
    "Collaborated with cross-functional teams to migrate legacy systems into modern AWS-based serverless architectures.",
    "Strong knowledge of containerization (Docker, Kubernetes) and hybrid cloud deployments across AWS and Azure.",
    "Skilled in test-driven development (TDD) with extensive use of JUnit, PyTest, Mocha, and integration testing frameworks."
    "Deep expertise in designing secure, scalable, and cost-optimized serverless applications leveraging AWS best practices and Well-Architected Framework.",
    "Architected event-driven solutions using AWS SNS, SQS, and EventBridge to enable loosely coupled, highly available microservices.",
    "Implemented multi-region, fault-tolerant architectures with automated failover and disaster recovery strategies.",
    "Led the development of automated CI/CD pipelines using Infrastructure as Code (IaC) with AWS CDK and CloudFormation, ensuring rapid and reliable deployments.",
    "Designed and maintained comprehensive logging, monitoring, and alerting systems using CloudWatch Logs, Metrics, and custom dashboards.",
    "Expertise in securing APIs with OAuth2, JWT, and AWS Cognito, including fine-grained access control and auditing.",
    "Experience with container orchestration and hybrid cloud deployments, integrating Kubernetes clusters with AWS EKS and Azure AKS environments.",
    "Proficient in performance tuning and cost optimization of serverless workloads, including Lambda cold start mitigation and efficient DynamoDB design.",
    "Strong collaborative skills working with DevOps, QA, and product teams to deliver enterprise-grade cloud solutions on aggressive timelines.",
    "Mentored junior developers and led technical workshops on AWS serverless best practices and microservices design patterns."
]
for point in summary_points:
    para = doc.add_paragraph(point, style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.25)

# --- Technical Skills ---
add_heading("Technical Skills")

table = doc.add_table(rows=6, cols=2)
table.style = "Table Grid"
for row in table.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.left_indent = Inches(0.25)

# Fill table
skills = [
    ("Languages", "Java, JavaScript, TypeScript, Python, Kotlin, SQL, Shell Scripting"),
    ("Frontend", "React.js, Angular, Vue.js, SCSS, Tailwind, Bootstrap, NativeScript"),
    ("Backend", "Java (Spring Boot, Spring Security, Hibernate), Node.js, Express, Kotlin"),
    ("Databases", "MongoDB, PostgreSQL, MySQL, Oracle, Redis, DynamoDB"),
    ("Cloud/DevOps", "AWS (Lambda, ECS, EC2, S3, CDK), Azure, Docker, Kubernetes, Jenkins, GitLab"),
    ("Messaging & CI/CD", "Kafka, RabbitMQ, Maven, Gradle, Git, GitHub Actions, GitLab CI/CD")
]
for i, (col1, col2) in enumerate(skills):
    table.cell(i,0).text = col1
    table.cell(i,1).text = col2

# --- Professional Experience ---
add_heading("Professional Experience")

def add_job(company, title, duration, responsibilities, environment):
    # Company
    comp = doc.add_paragraph()
    comp.paragraph_format.left_indent = Inches(0.25)
    run = comp.add_run(f"{company} – {title} ({duration})")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)

    # Responsibilities
    resp = doc.add_paragraph()
    resp.paragraph_format.left_indent = Inches(0.5)
    r = resp.add_run("Responsibilities")
    r.bold = True

    for res in responsibilities:
        p = doc.add_paragraph(res, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.75)

    # Environment
    env = doc.add_paragraph()
    env.paragraph_format.left_indent = Inches(0.5)
    env.add_run(f"Environment: {environment}").italic = True

# Example usage for Boeing
boeing_resp = [
    "Designed and implemented enterprise-grade full stack applications using Java 17, Spring Boot, and React.js for real-time aircraft analytics.",
    "Built secure RESTful APIs and microservices leveraging Spring Security and OAuth2.",
    "Integrated Apache Kafka for real-time data streaming and processing.",
    "Engineered asynchronous, fault-tolerant workflows using RabbitMQ.",
    "Developed cloud-native services deployed to Azure Kubernetes Service and App Services.",
    "Utilized Azure Data Lake, Blob Storage, and Cosmos DB for large-scale datasets.",
    "Collaborated with UI/UX designers to create responsive web interfaces using React.js and Tailwind.",
    "Automated test cases using JUnit, Mockito, and Selenium.",
    "Implemented CI/CD pipelines using Azure DevOps with unit testing and automated deployments.",
    "Monitored application health and performance using Azure Monitor and Application Insights.",
    "Enforced security best practices including role-based access control and data encryption at rest and in transit.",
    "Worked closely with cross-functional teams including QA, DevOps, and product management to deliver features on schedule.",
    "Developed comprehensive testing strategies including integration, load, and end-to-end tests to ensure reliability."]

boeing_env = "Java 17, Spring Boot, React.js, TypeScript, Tailwind, OAuth2, Kafka, RabbitMQ, Azure AKS, Cosmos DB, Data Lake, JUnit, Mockito, Selenium, Azure DevOps"
add_job("Boeing", "Sr. Java Full Stack Developer", "Jan 2024 – Present", boeing_resp, boeing_env)

# Amazon Web Services Experience
aws_resp = [
    "Architected and implemented serverless applications using AWS Lambda, API Gateway, and DynamoDB to support high-scale customer-facing services.",
    "Designed event-driven microservices architectures leveraging SNS, SQS, and EventBridge for decoupled and resilient workflows.",
    "Developed Infrastructure as Code using AWS CDK and CloudFormation to automate cloud resource provisioning and configuration.",
    "Implemented CI/CD pipelines with AWS CodePipeline, CodeBuild, and CodeDeploy for continuous integration and delivery.",
    "Configured monitoring and alerting solutions using CloudWatch Logs, Metrics, and custom dashboards for operational visibility.",
    "Applied security best practices including IAM roles, policies, and AWS Cognito for authentication and authorization.",
    "Optimized Lambda functions for performance and cost by tuning memory, timeout settings, and cold start mitigation techniques.",
    "Collaborated with cross-team stakeholders to design scalable APIs and backend systems aligned with business requirements.",
    "Led migration efforts from monolithic applications to microservices-based architectures on AWS.",
    "Provided technical leadership and mentorship for junior engineers on AWS best practices and serverless design."
]
aws_env = "AWS Lambda, API Gateway, DynamoDB, SNS, SQS, EventBridge, AWS CDK, CloudFormation, CodePipeline, CodeBuild, CodeDeploy, CloudWatch, IAM, Cognito"
add_job("Amazon Web Services", "AWS Serverless Solutions Architect", "Jun 2021 – Dec 2023", aws_resp, aws_env)

# Capgemini Experience
capgemini_resp = [
    "Developed and maintained microservices using Spring Boot and Node.js deployed on AWS ECS and EKS clusters.",
    "Designed and implemented scalable RESTful APIs and GraphQL endpoints for enterprise clients.",
    "Managed database schema design and optimization for PostgreSQL and MongoDB to support high throughput applications.",
    "Automated build, test, and deployment pipelines using Jenkins and GitLab CI/CD.",
    "Implemented monitoring and logging using ELK stack and Prometheus to ensure system reliability and performance.",
    "Collaborated with security teams to enforce compliance with GDPR and other regulatory requirements.",
    "Led agile ceremonies and coordinated efforts between distributed teams across multiple time zones.",
    "Conducted code reviews and established coding standards to improve code quality and maintainability.",
    "Integrated third-party services and APIs to extend application functionality and improve user experience.",
    "Provided technical training and documentation to support knowledge sharing within the team."
]
capgemini_env = "Spring Boot, Node.js, AWS ECS, EKS, PostgreSQL, MongoDB, Jenkins, GitLab CI/CD, ELK, Prometheus, Agile"
add_job("Capgemini", "Senior Cloud Developer", "Mar 2018 – May 2021", capgemini_resp, capgemini_env)

# Godrej Experience
godrej_resp = [
    "Led development of cloud-native applications using Java, Kotlin, and AWS serverless technologies.",
    "Designed event-driven systems with AWS Step Functions orchestrating Lambda-based workflows.",
    "Implemented data pipelines and ETL processes using AWS Glue and Lambda for large-scale data processing.",
    "Developed and maintained CI/CD pipelines utilizing GitHub Actions and Jenkins for automated deployments.",
    "Collaborated with product owners and business analysts to translate requirements into technical solutions.",
    "Optimized application performance and scalability through profiling and resource tuning.",
    "Ensured application security by integrating AWS WAF, Shield, and secure coding practices.",
    "Established monitoring and alerting using CloudWatch and custom metrics to proactively address issues.",
    "Conducted root cause analysis and troubleshooting for production incidents and outages.",
    "Mentored junior developers and promoted best practices in cloud architecture and software development."
]
godrej_env = "Java, Kotlin, AWS Lambda, Step Functions, Glue, GitHub Actions, Jenkins, AWS WAF, Shield, CloudWatch"
add_job("Godrej", "Cloud Solutions Engineer", "Jul 2016 – Feb 2018", godrej_resp, godrej_env)


# --- Education ---
add_heading("Education")
edu1 = doc.add_paragraph("Masters in Data Science, University at Buffalo (2021 – 2022)")
edu2 = doc.add_paragraph("Bachelors in Computer Science, Sreenidhi Institute of Technology (2012 – 2016)")
edu1.paragraph_format.left_indent = Inches(0.25)
edu2.paragraph_format.left_indent = Inches(0.25)

# Save
doc.save("Nathaniel_Jason_AWS_Developer_Contract.docx")